import PyPDF2
import docx
import xlrd
import re
import math
import csv
from pdfrw import PdfReader


def apply(file_path):
    """Method called on the path that choose which function to use in function of the type """
    file_type = get_file_type(file_path)
    if file_type == 'EXCEL':
            return exctract_excel(file_path)
    if file_type == 'PDF':
            return exctract_pdf(file_path)
    if file_type == 'WORD':
            return exctract_word(file_path)

        
def mapperOpener(path):
    """ Use to open the mapper ID-path"""
    mapping = []
    with open(path, 'r') as csvfile:
        spamreader = csv.reader(csvfile,delimiter='|')
        for row in spamreader:
            mapping.append((row[0],row[1],row[2]))
    return mapping

def get_file_type(file_path):
    """ Return which type of document we are given in argument"""
    if file_path[-3:] == 'pdf':
        return 'PDF'
    if file_path[-3:] == 'xls' or file_path[-4:] == 'xlsx':
        return 'EXCEL'
    if file_path[-3:] == 'doc' or file_path[-4:] == 'docx':
        return 'WORD'
    raise Exception('Unexpected file type')


def exctract_excel(path):
    """ Method to extract an Excel.
    It basically concatenates all row together
    Return : All row values concatenated into a string.
    """
    strRe = ""
    try:
        wb = xlrd.open_workbook(path)
        for j in range(wb.nsheets):
            sh1 = wb.sheet_by_index(j)
            for rownum in range(sh1.nrows): # sh1.nrows -> number of rows (ncols -> num columns) 
                row = sh1.row_values(rownum)
                for i in row :
                    strRe += str(i) + " "
                strRe+="\n"
    except:
        print("Could not open Excel")
    return strRe
    
def exctract_pdf(path):
    """
    PDF parser
    Return:  a string containing text from a pdf
    """
    try:
        pdf = PyPDF2.PdfFileReader(path)
    except:
        pdf = None
    text = ''
    if(pdf is not None):
        try:
            for i in range(pdf.getNumPages()):
                page = pdf.getPage(i)
                text+= ' ' + page.extractText()
        except:
            text = ' '
    else:
        print("Could not parse the PDF, ", path)
    return text 

def exctract_word(path):
    """
    Extracter of word
    If its is a .doc, take the txt equivalent that we transformed with Antiword
    Return : string containg all word text
    """
    if(path.endswith("doc")):
        txt = path[:-3]+'txt'
        f = open(txt,encoding="utf8")
        result = f.read()
    else :
        try:
            doc = docx.Document(path)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            result = '\n'.join(fullText)
        except:
            result = ""
    return result

#Unused
def count_occurence(text,company):
    """
    Return : number of occurence of the word given in argument in text
    """
    return len(re.findall(company.lower(),text.lower()))

#Unused
def frequency_occurence(text,company):
    """
    Return : A frequency count of the word given, log normalized, in text
    """
    try:
        countO = count_occurence(text,company)
    except:
        countO = 0
    if(len(text)>0 and countO!=0 ):
        return countO/math.log(len(text.split()))
    else:
        return 0

#Unused
def extract_title(path):
    """Extract title from a given document (not really useful)"""
    if path[-3:] == 'pdf':
        try:
            pdf = PdfReader(path).Info.Title
            return pdf
        except:
            return " "
    if path[-3:] == 'xls' or path[-4:] == 'xlsx':
        try:
            return " "
        except:
            return " " 
    if path[-3:] == 'doc' or path[-4:] == 'docx':
        try:
            doc = docx.Document(path)
            return doc.core_properties.title
        except:
            return " "
    raise Exception('Unexpected file type')